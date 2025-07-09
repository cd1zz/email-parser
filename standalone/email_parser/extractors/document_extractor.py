# ============================================================================
# email_parser/extractors/document_extractor.py
# ============================================================================
"""Document text extraction for Office documents and PDFs with comprehensive Excel URL detection."""

import io
import logging
import zipfile
import xml.etree.ElementTree as ET
import re
from typing import Optional, Dict, Any, List
from dataclasses import dataclass


@dataclass
class DocumentExtractionResult:
    """Result of document text extraction."""
    text_content: Optional[str] = None
    success: bool = False
    error_message: Optional[str] = None
    document_type: Optional[str] = None
    extraction_method: Optional[str] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentTextExtractor:
    """Extracts text content from various document types with comprehensive URL detection."""
    
    def __init__(self, logger: logging.Logger):
        self.logger = logger
    
    def extract_text(self, data: bytes, filename: str = None, content_type: str = None) -> DocumentExtractionResult:
        """Extract text from document data based on type detection."""
        if not data:
            return DocumentExtractionResult(
                success=False,
                error_message="No data provided"
            )
        
        # Determine document type
        doc_type = self._detect_document_type(data, filename, content_type)
        
        self.logger.debug(f"Detected document type: {doc_type} for file: {filename}")
        
        if doc_type == 'pdf':
            return self._extract_pdf_text(data)
        elif doc_type in ['xlsx', 'xls']:
            return self._extract_excel_text(data, filename)
        elif doc_type in ['docx', 'doc']:
            return self._extract_word_text(data)
        else:
            return DocumentExtractionResult(
                success=False,
                error_message=f"Unsupported document type: {doc_type}",
                document_type=doc_type
            )
    
    def _detect_document_type(self, data: bytes, filename: str = None, content_type: str = None) -> str:
        """Detect document type from magic bytes, filename, and content type."""
        
        # Check magic bytes first (most reliable)
        if len(data) >= 8:
            # PDF magic bytes
            if data.startswith(b'%PDF-'):
                return 'pdf'
            
            # Office documents (OLE format) - includes .doc, .xls, .msg
            if data.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                # Could be .doc, .xls, or .msg - need filename to distinguish
                if filename:
                    filename_lower = filename.lower()
                    if filename_lower.endswith('.doc'):
                        return 'doc'
                    elif filename_lower.endswith('.xls'):
                        return 'xls'
                # Default to doc for OLE documents
                return 'doc'
            
            # Modern Office documents (ZIP-based)
            if data.startswith(b'PK\x03\x04'):
                if filename:
                    filename_lower = filename.lower()
                    if filename_lower.endswith('.docx'):
                        return 'docx'
                    elif filename_lower.endswith('.xlsx'):
                        return 'xlsx'
                # Try to detect from content
                return self._detect_office_from_zip_content(data)
        
        # Fall back to filename extension
        if filename:
            filename_lower = filename.lower()
            if filename_lower.endswith('.pdf'):
                return 'pdf'
            elif filename_lower.endswith(('.doc', '.docx')):
                return 'docx' if filename_lower.endswith('.docx') else 'doc'
            elif filename_lower.endswith(('.xls', '.xlsx')):
                return 'xlsx' if filename_lower.endswith('.xlsx') else 'xls'
        
        # Fall back to content type
        if content_type:
            content_type_lower = content_type.lower()
            if 'pdf' in content_type_lower:
                return 'pdf'
            elif 'wordprocessingml' in content_type_lower or 'msword' in content_type_lower:
                return 'docx' if 'openxml' in content_type_lower else 'doc'
            elif 'spreadsheetml' in content_type_lower or 'excel' in content_type_lower:
                return 'xlsx' if 'openxml' in content_type_lower else 'xls'
        
        return 'unknown'
    
    def _detect_office_from_zip_content(self, data: bytes) -> str:
        """Detect specific Office document type from ZIP content."""
        try:
            # Look for Office-specific files in the ZIP
            content_str = data[:2048].decode('latin-1', errors='ignore')
            
            if 'word/' in content_str or 'document.xml' in content_str:
                return 'docx'
            elif 'xl/' in content_str or 'workbook.xml' in content_str:
                return 'xlsx'
            else:
                # Default to docx for unknown ZIP-based Office docs
                return 'docx'
        except Exception:
            return 'docx'
    
    def _extract_excel_text(self, excel_data: bytes, filename: str = None) -> DocumentExtractionResult:
        """COMPREHENSIVE: Excel text extraction that finds URLs in all parts of the file."""
        if not excel_data:
            return DocumentExtractionResult(
                success=False,
                error_message="Empty Excel file",
                document_type='excel'
            )
        
        self.logger.info(f"Starting comprehensive Excel extraction for {filename}, data size: {len(excel_data)} bytes")
        
        # Collection of all content and URLs
        all_text_content = []
        all_urls = set()
        extraction_methods_used = []
        
        # Method 1: Try pandas for worksheet data
        pandas_success = False
        try:
            import pandas as pd
            self.logger.debug("Trying pandas extraction...")
            
            excel_file = io.BytesIO(excel_data)
            
            # Try different pandas approaches
            approaches = [
                {"engine": "openpyxl"},
                {"engine": "openpyxl", "header": None},
                {"engine": "openpyxl", "skiprows": 0, "header": None},
            ]
            
            # Add xlrd if available
            if self._has_xlrd():
                approaches.append({"engine": "xlrd"})
            
            for approach in approaches:
                try:
                    excel_file.seek(0)
                    dfs = pd.read_excel(excel_file, sheet_name=None, **approach)
                    
                    text_parts = []
                    for sheet_name, df in dfs.items():
                        if not df.empty:
                            non_null_count = df.count().sum()
                            if non_null_count > 0:
                                text_parts.append(f"Sheet: {sheet_name}")
                                df_text = df.fillna('').to_string(index=False)
                                text_parts.append(df_text)
                                text_parts.append("")
                    
                    if text_parts:
                        pandas_text = "\n".join(text_parts)
                        all_text_content.append(pandas_text)
                        all_text_content.append("")
                        extraction_methods_used.append(f"pandas_{approach.get('engine', 'default')}")
                        pandas_success = True
                        self.logger.info(f"✓ Pandas extraction successful with {approach}")
                        break
                        
                except Exception as e:
                    self.logger.debug(f"Pandas approach {approach} failed: {e}")
                    continue
            
        except ImportError:
            self.logger.debug("pandas not available")
        except Exception as e:
            self.logger.debug(f"Pandas extraction failed: {e}")
        
        # Method 2: Comprehensive ZIP analysis (THIS IS THE KEY FIX FOR FINDING URLs)
        zip_success = False
        try:
            self.logger.info("Starting comprehensive ZIP structure analysis...")
            
            excel_file = io.BytesIO(excel_data)
            
            with zipfile.ZipFile(excel_file, 'r') as zip_file:
                # Files that commonly contain URLs and external references
                priority_files = [
                    '_rels/.rels',
                    'xl/_rels/workbook.xml.rels',
                    'xl/worksheets/_rels/sheet1.xml.rels',
                    'xl/drawings/_rels/drawing1.xml.rels',  # This is where OneDrive URLs hide!
                    'xl/drawings/drawing1.xml',
                    'xl/worksheets/sheet1.xml',
                    'xl/sharedStrings.xml',
                    'xl/comments1.xml',
                ]
                
                # Also check all .rels files and drawing files
                available_files = zip_file.namelist()
                for file_path in available_files:
                    if (file_path.endswith('.rels') or 
                        'drawings' in file_path or 
                        'comments' in file_path or
                        'worksheets' in file_path):
                        if file_path not in priority_files:
                            priority_files.append(file_path)
                
                self.logger.info(f"Analyzing {len(priority_files)} files for content and URLs")
                
                zip_text_parts = []
                
                for file_path in priority_files:
                    if file_path in available_files:
                        try:
                            content = zip_file.read(file_path)
                            
                            # Decode content
                            try:
                                text_content = content.decode('utf-8')
                            except UnicodeDecodeError:
                                text_content = content.decode('latin-1', errors='ignore')
                            
                            # Extract URLs from this file
                            file_urls = self._extract_urls_from_text(text_content)
                            if file_urls:
                                # Filter out XML schema URLs
                                real_urls = [url for url in file_urls if not self._is_schema_url(url)]
                                if real_urls:
                                    all_urls.update(real_urls)
                                    self.logger.info(f"Found {len(real_urls)} URLs in {file_path}: {real_urls}")
                            
                            # Extract meaningful text (not just XML schemas)
                            meaningful_text = self._extract_meaningful_content(text_content, file_path)
                            if meaningful_text:
                                zip_text_parts.append(f"{file_path}:")
                                zip_text_parts.append(meaningful_text)
                                zip_text_parts.append("")
                            
                        except Exception as e:
                            self.logger.debug(f"Error processing {file_path}: {e}")
                
                if zip_text_parts:
                    all_text_content.extend(zip_text_parts)
                    zip_success = True
                    extraction_methods_used.append("zip_analysis")
                    self.logger.info("✓ ZIP analysis found additional content")
        
        except Exception as e:
            self.logger.warning(f"ZIP analysis failed: {e}")
        
        # Method 3: Raw URL search as fallback
        try:
            raw_text = excel_data.decode('utf-8', errors='ignore')
            raw_urls = self._extract_urls_from_text(raw_text)
            if raw_urls:
                real_raw_urls = [url for url in raw_urls if not self._is_schema_url(url)]
                if real_raw_urls:
                    all_urls.update(real_raw_urls)
                    extraction_methods_used.append("raw_search")
                    self.logger.info(f"Raw search found {len(real_raw_urls)} additional URLs")
        except Exception as e:
            self.logger.debug(f"Raw URL search failed: {e}")
        
        # Compile final results
        final_text = "\n".join(all_text_content).strip()
        final_urls = list(all_urls)
        
        # Determine success
        if final_text or final_urls:
            success_text = final_text if final_text else "[No text content found, but URLs extracted]"
            
            return DocumentExtractionResult(
                text_content=success_text,
                success=True,
                document_type='excel',
                extraction_method="+".join(extraction_methods_used) if extraction_methods_used else "comprehensive",
                metadata={
                    'character_count': len(final_text),
                    'urls_found': final_urls,
                    'url_count': len(final_urls),
                    'pandas_success': pandas_success,
                    'zip_analysis_success': zip_success,
                    'methods_used': extraction_methods_used
                }
            )
        else:
            return DocumentExtractionResult(
                success=False,
                error_message="No content or URLs found with comprehensive extraction",
                document_type='excel',
                metadata={
                    'methods_attempted': extraction_methods_used,
                    'pandas_success': pandas_success,
                    'zip_analysis_success': zip_success
                }
            )
    
    def _extract_urls_from_text(self, text: str) -> List[str]:
        """Extract URLs from text using comprehensive patterns."""
        url_patterns = [
            r'https?://[^\s<>"{}|\\^`\[\]\']+',  # HTTP/HTTPS URLs
            r'ftp://[^\s<>"{}|\\^`\[\]\']+',     # FTP URLs
            r'www\.[^\s<>"{}|\\^`\[\]\']+',      # www URLs without protocol
        ]
        
        urls = []
        for pattern in url_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            urls.extend(matches)
        
        return urls
    
    def _is_schema_url(self, url: str) -> bool:
        """Check if URL is an XML schema URL (not a real external link)."""
        schema_patterns = [
            'http://schemas.',
            'http://www.w3.org/',
            'http://purl.org/',
            'http://ns.adobe.com/',
        ]
        
        return any(url.startswith(pattern) for pattern in schema_patterns)
    
    def _extract_meaningful_content(self, xml_content: str, file_path: str) -> str:
        """Extract meaningful content from XML, filtering out schema noise."""
        try:
            # Parse as XML
            root = ET.fromstring(xml_content)
            
            meaningful_parts = []
            
            # Extract text content and meaningful attributes
            for elem in root.iter():
                # Check element text
                if elem.text and elem.text.strip():
                    text = elem.text.strip()
                    if (not self._is_schema_url(text) and 
                        len(text) > 2 and 
                        not text.startswith('<?xml')):
                        meaningful_parts.append(text)
                
                # Check attributes for meaningful content (like URLs)
                for attr_name, attr_value in elem.attrib.items():
                    if (attr_value and 
                        ('http://' in attr_value or 'https://' in attr_value or 'www.' in attr_value) and
                        not self._is_schema_url(attr_value)):
                        meaningful_parts.append(f"{attr_name}: {attr_value}")
            
            if meaningful_parts:
                return "\n".join(meaningful_parts[:10])  # Limit output
                
        except ET.ParseError:
            # If not valid XML, extract non-XML lines
            lines = [line.strip() for line in xml_content.split('\n')]
            meaningful_lines = [
                line for line in lines 
                if (line and 
                    not line.startswith('<') and 
                    not line.startswith('<?xml') and
                    not self._is_schema_url(line) and
                    len(line) > 5)
            ]
            if meaningful_lines:
                return "\n".join(meaningful_lines[:5])
        
        except Exception as e:
            self.logger.debug(f"Error extracting meaningful content from {file_path}: {e}")
        
        return ""
    
    def _has_xlrd(self) -> bool:
        """Check if xlrd is available."""
        try:
            import xlrd
            return True
        except ImportError:
            return False
    
    def _extract_pdf_text(self, pdf_data: bytes) -> DocumentExtractionResult:
        """Extract text from PDF using pdfminer."""
        try:
            # Try pdfminer first (most reliable)
            try:
                from pdfminer.high_level import extract_text
                
                pdf_file = io.BytesIO(pdf_data)
                text = extract_text(pdf_file)
                
                if text and text.strip():
                    return DocumentExtractionResult(
                        text_content=text.strip(),
                        success=True,
                        document_type='pdf',
                        extraction_method='pdfminer',
                        metadata={'character_count': len(text.strip())}
                    )
                else:
                    return DocumentExtractionResult(
                        success=False,
                        error_message="PDF contains no extractable text",
                        document_type='pdf'
                    )
                    
            except ImportError:
                self.logger.warning("pdfminer not available for PDF text extraction")
                return DocumentExtractionResult(
                    success=False,
                    error_message="pdfminer library not available",
                    document_type='pdf'
                )
                
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF: {e}")
            return DocumentExtractionResult(
                success=False,
                error_message=f"PDF extraction failed: {str(e)}",
                document_type='pdf'
            )
    
    def _extract_word_text(self, word_data: bytes) -> DocumentExtractionResult:
        """Extract text from Word documents using python-docx and fallbacks."""
        
        # Try python-docx for .docx files
        try:
            from docx import Document
            
            word_file = io.BytesIO(word_data)
            doc = Document(word_file)
            
            text_content = []
            paragraph_count = 0
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    paragraph_count += 1
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                text_content.append("\n")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
                table_count += 1
            
            if text_content:
                final_text = "\n".join(text_content).strip()
                return DocumentExtractionResult(
                    text_content=final_text,
                    success=True,
                    document_type='word',
                    extraction_method='python-docx',
                    metadata={
                        'paragraph_count': paragraph_count,
                        'table_count': table_count,
                        'character_count': len(final_text)
                    }
                )
            else:
                return DocumentExtractionResult(
                    success=False,
                    error_message="Word document contains no extractable text",
                    document_type='word'
                )
                
        except ImportError:
            self.logger.warning("python-docx not available, trying textract fallback")
            return self._extract_word_fallback(word_data)
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document with python-docx: {e}")
            return self._extract_word_fallback(word_data)
    
    def _extract_word_fallback(self, word_data: bytes) -> DocumentExtractionResult:
        """Fallback method for Word document text extraction."""
        # Try textract
        try:
            import textract
            
            # textract expects a file path, so we need to write to temp file
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(word_data)
                tmp_file.flush()
                tmp_path = tmp_file.name
            
            try:
                text = textract.process(tmp_path).decode('utf-8')
                
                if text and text.strip():
                    return DocumentExtractionResult(
                        text_content=text.strip(),
                        success=True,
                        document_type='word',
                        extraction_method='textract',
                        metadata={'character_count': len(text.strip())}
                    )
                else:
                    return DocumentExtractionResult(
                        success=False,
                        error_message="Word document contains no extractable text",
                        document_type='word'
                    )
            finally:
                # Clean up temp file
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
                    
        except ImportError:
            self.logger.warning("textract not available for Word processing")
            return DocumentExtractionResult(
                success=False,
                error_message="No suitable library available for Word document processing (python-docx or textract required)",
                document_type='word'
            )
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document with textract: {e}")
            return DocumentExtractionResult(
                success=False,
                error_message=f"Word extraction failed: {str(e)}",
                document_type='word'
            )


class DocumentProcessor:
    """High-level processor that integrates document extraction with email parsing."""
    
    def __init__(self, logger: logging.Logger, url_analyzer=None):
        self.logger = logger
        self.document_extractor = DocumentTextExtractor(logger)
        self.url_analyzer = url_analyzer
    
    def process_document_attachment(self, attachment_data: bytes, filename: str = None, 
                                  content_type: str = None) -> Dict[str, Any]:
        """Process a document attachment and extract text and URLs."""
        result = {
            'extraction_result': None,
            'extracted_text': None,
            'urls_found': [],
            'processing_success': False
        }
        
        try:
            # Extract text from document
            extraction_result = self.document_extractor.extract_text(
                attachment_data, filename, content_type
            )
            
            result['extraction_result'] = {
                'success': extraction_result.success,
                'document_type': extraction_result.document_type,
                'extraction_method': extraction_result.extraction_method,
                'error_message': extraction_result.error_message,
                'metadata': extraction_result.metadata
            }
            
            if extraction_result.success and extraction_result.text_content:
                result['extracted_text'] = extraction_result.text_content
                result['processing_success'] = True
                
                # Get URLs from metadata if available (for Excel comprehensive extraction)
                if (extraction_result.metadata and 
                    'urls_found' in extraction_result.metadata):
                    result['urls_found'] = extraction_result.metadata['urls_found']
                    self.logger.info(f"Found {len(result['urls_found'])} URLs in document metadata")
                
                # Also extract URLs from the document text if URL analyzer available
                if self.url_analyzer:
                    text_urls = self._extract_urls_from_text(extraction_result.text_content)
                    # Combine with metadata URLs, removing duplicates
                    all_urls = list(set(result['urls_found'] + text_urls))
                    result['urls_found'] = all_urls
                    
                self.logger.info(f"Successfully processed document {filename}: "
                               f"{len(extraction_result.text_content)} chars, "
                               f"{len(result['urls_found'])} URLs")
            else:
                self.logger.warning(f"Failed to extract text from document {filename}: "
                                  f"{extraction_result.error_message}")
        
        except Exception as e:
            self.logger.error(f"Error processing document attachment {filename}: {e}")
            result['extraction_result'] = {
                'success': False,
                'error_message': f"Processing failed: {str(e)}"
            }
        
        return result
    
    def _extract_urls_from_text(self, text: str) -> List[str]:
        """Extract URLs from text content using URL analyzer."""
        try:
            # Create a temporary structure for URL analysis
            temp_structure = {
                'body': {
                    'text': text,
                    'html': None
                }
            }
            
            # Use the URL analyzer to extract URLs
            if hasattr(self.url_analyzer, 'analyze_email_urls'):
                analysis = self.url_analyzer.analyze_email_urls(temp_structure)
                return analysis.final_urls
            else:
                # Fallback to simple regex extraction
                url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
                urls = re.findall(url_pattern, text)
                return list(set(urls))  # Remove duplicates
                
        except Exception as e:
            self.logger.error(f"Error extracting URLs from document text: {e}")
            return []