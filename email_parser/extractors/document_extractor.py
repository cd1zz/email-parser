# ============================================================================
# email_parser/extractors/document_extractor.py
# ============================================================================
"""Document text extraction for Office documents and PDFs."""

import io
import logging
from typing import Optional, Dict, Any, List
from dataclasses import dataclass


@dataclass
class DocumentExtractionResult:
    """Result of document text extraction."""
    text_content: Optional[str] = None
    success: bool = False
    error_message: Optional[str] = None
    document_type: Optional[str] = None
    extraction_method: Optional[str] = None
    metadata: Dict[str, Any] = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


class DocumentTextExtractor:
    """Extracts text content from various document types."""
    
    def __init__(self, logger: logging.Logger):
        self.logger = logger
    
    def extract_text(self, data: bytes, filename: str = None, content_type: str = None) -> DocumentExtractionResult:
        """Extract text from document data based on type detection."""
        if not data:
            return DocumentExtractionResult(
                success=False,
                error_message="No data provided"
            )
        
        # Determine document type
        doc_type = self._detect_document_type(data, filename, content_type)
        
        self.logger.debug(f"Detected document type: {doc_type} for file: {filename}")
        
        if doc_type == 'pdf':
            return self._extract_pdf_text(data)
        elif doc_type in ['xlsx', 'xls']:
            return self._extract_excel_text(data)
        elif doc_type in ['docx', 'doc']:
            return self._extract_word_text(data)
        else:
            return DocumentExtractionResult(
                success=False,
                error_message=f"Unsupported document type: {doc_type}",
                document_type=doc_type
            )
    
    def _detect_document_type(self, data: bytes, filename: str = None, content_type: str = None) -> str:
        """Detect document type from magic bytes, filename, and content type."""
        
        # Check magic bytes first (most reliable)
        if len(data) >= 8:
            # PDF magic bytes
            if data.startswith(b'%PDF-'):
                return 'pdf'
            
            # Office documents (OLE format) - includes .doc, .xls, .msg
            if data.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
                # Could be .doc, .xls, or .msg - need filename to distinguish
                if filename:
                    filename_lower = filename.lower()
                    if filename_lower.endswith('.doc'):
                        return 'doc'
                    elif filename_lower.endswith('.xls'):
                        return 'xls'
                # Default to doc for OLE documents
                return 'doc'
            
            # Modern Office documents (ZIP-based)
            if data.startswith(b'PK\x03\x04'):
                if filename:
                    filename_lower = filename.lower()
                    if filename_lower.endswith('.docx'):
                        return 'docx'
                    elif filename_lower.endswith('.xlsx'):
                        return 'xlsx'
                # Try to detect from content
                return self._detect_office_from_zip_content(data)
        
        # Fall back to filename extension
        if filename:
            filename_lower = filename.lower()
            if filename_lower.endswith('.pdf'):
                return 'pdf'
            elif filename_lower.endswith(('.doc', '.docx')):
                return 'docx' if filename_lower.endswith('.docx') else 'doc'
            elif filename_lower.endswith(('.xls', '.xlsx')):
                return 'xlsx' if filename_lower.endswith('.xlsx') else 'xls'
        
        # Fall back to content type
        if content_type:
            content_type_lower = content_type.lower()
            if 'pdf' in content_type_lower:
                return 'pdf'
            elif 'wordprocessingml' in content_type_lower or 'msword' in content_type_lower:
                return 'docx' if 'openxml' in content_type_lower else 'doc'
            elif 'spreadsheetml' in content_type_lower or 'excel' in content_type_lower:
                return 'xlsx' if 'openxml' in content_type_lower else 'xls'
        
        return 'unknown'
    
    def _detect_office_from_zip_content(self, data: bytes) -> str:
        """Detect specific Office document type from ZIP content."""
        try:
            # Look for Office-specific files in the ZIP
            content_str = data[:2048].decode('latin-1', errors='ignore')
            
            if 'word/' in content_str or 'document.xml' in content_str:
                return 'docx'
            elif 'xl/' in content_str or 'workbook.xml' in content_str:
                return 'xlsx'
            else:
                # Default to docx for unknown ZIP-based Office docs
                return 'docx'
        except Exception:
            return 'docx'
    
    def _extract_pdf_text(self, pdf_data: bytes) -> DocumentExtractionResult:
        """Extract text from PDF using pdfminer."""
        try:
            # Try pdfminer first (most reliable)
            try:
                from pdfminer.high_level import extract_text
                
                pdf_file = io.BytesIO(pdf_data)
                text = extract_text(pdf_file)
                
                if text and text.strip():
                    return DocumentExtractionResult(
                        text_content=text.strip(),
                        success=True,
                        document_type='pdf',
                        extraction_method='pdfminer',
                        metadata={'character_count': len(text.strip())}
                    )
                else:
                    return DocumentExtractionResult(
                        success=False,
                        error_message="PDF contains no extractable text",
                        document_type='pdf'
                    )
                    
            except ImportError:
                self.logger.warning("pdfminer not available for PDF text extraction")
                return DocumentExtractionResult(
                    success=False,
                    error_message="pdfminer library not available",
                    document_type='pdf'
                )
                
        except Exception as e:
            self.logger.error(f"Error extracting text from PDF: {e}")
            return DocumentExtractionResult(
                success=False,
                error_message=f"PDF extraction failed: {str(e)}",
                document_type='pdf'
            )
    
    def _extract_excel_text(self, excel_data: bytes) -> DocumentExtractionResult:
        """Extract text from Excel files using pandas with multiple fallbacks."""
        if not excel_data:
            return DocumentExtractionResult(
                success=False,
                error_message="Empty Excel file",
                document_type='excel'
            )
        
        # Try pandas first (most comprehensive)
        try:
            import pandas as pd
            
            excel_file = io.BytesIO(excel_data)
            
            # Read all sheets
            try:
                # For .xlsx files
                dfs = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
            except Exception:
                try:
                    # For .xls files  
                    excel_file.seek(0)
                    dfs = pd.read_excel(excel_file, sheet_name=None, engine='xlrd')
                except Exception:
                    # Generic fallback
                    excel_file.seek(0)
                    dfs = pd.read_excel(excel_file, sheet_name=None)
            
            text_content = []
            total_cells = 0
            
            for sheet_name, df in dfs.items():
                if not df.empty:
                    # Add sheet header
                    text_content.append(f"=== Sheet: {sheet_name} ===")
                    
                    # Convert dataframe to text, handling NaN values
                    df_filled = df.fillna('')
                    df_text = df_filled.to_string(index=False)
                    text_content.append(df_text)
                    text_content.append("")  # Empty line between sheets
                    
                    total_cells += df.size
            
            if text_content:
                final_text = "\n".join(text_content).strip()
                return DocumentExtractionResult(
                    text_content=final_text,
                    success=True,
                    document_type='excel',
                    extraction_method='pandas',
                    metadata={
                        'sheet_count': len(dfs),
                        'total_cells': total_cells,
                        'character_count': len(final_text)
                    }
                )
            else:
                return DocumentExtractionResult(
                    success=False,
                    error_message="Excel file contains no data",
                    document_type='excel'
                )
                
        except ImportError:
            self.logger.warning("pandas not available for Excel processing")
            return DocumentExtractionResult(
                success=False,
                error_message="pandas library not available for Excel processing",
                document_type='excel'
            )
        except Exception as e:
            self.logger.error(f"Error extracting text from Excel: {e}")
            return DocumentExtractionResult(
                success=False,
                error_message=f"Excel extraction failed: {str(e)}",
                document_type='excel'
            )
    
    def _extract_word_text(self, word_data: bytes) -> DocumentExtractionResult:
        """Extract text from Word documents using python-docx and fallbacks."""
        
        # Try python-docx for .docx files
        try:
            from docx import Document
            
            word_file = io.BytesIO(word_data)
            doc = Document(word_file)
            
            text_content = []
            paragraph_count = 0
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
                    paragraph_count += 1
            
            # Extract tables
            table_count = 0
            for table in doc.tables:
                text_content.append("\n=== TABLE ===")
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
                table_count += 1
            
            if text_content:
                final_text = "\n".join(text_content).strip()
                return DocumentExtractionResult(
                    text_content=final_text,
                    success=True,
                    document_type='word',
                    extraction_method='python-docx',
                    metadata={
                        'paragraph_count': paragraph_count,
                        'table_count': table_count,
                        'character_count': len(final_text)
                    }
                )
            else:
                return DocumentExtractionResult(
                    success=False,
                    error_message="Word document contains no extractable text",
                    document_type='word'
                )
                
        except ImportError:
            self.logger.warning("python-docx not available, trying textract fallback")
            return self._extract_word_fallback(word_data)
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document with python-docx: {e}")
            return self._extract_word_fallback(word_data)
    
    def _extract_word_fallback(self, word_data: bytes) -> DocumentExtractionResult:
        """Fallback method for Word document text extraction."""
        # Try textract
        try:
            import textract
            
            # textract expects a file path, so we need to write to temp file
            import tempfile
            import os
            
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(word_data)
                tmp_file.flush()
                tmp_path = tmp_file.name
            
            try:
                text = textract.process(tmp_path).decode('utf-8')
                
                if text and text.strip():
                    return DocumentExtractionResult(
                        text_content=text.strip(),
                        success=True,
                        document_type='word',
                        extraction_method='textract',
                        metadata={'character_count': len(text.strip())}
                    )
                else:
                    return DocumentExtractionResult(
                        success=False,
                        error_message="Word document contains no extractable text",
                        document_type='word'
                    )
            finally:
                # Clean up temp file
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass
                    
        except ImportError:
            self.logger.warning("textract not available for Word processing")
            return DocumentExtractionResult(
                success=False,
                error_message="No suitable library available for Word document processing (python-docx or textract required)",
                document_type='word'
            )
        except Exception as e:
            self.logger.error(f"Error extracting text from Word document with textract: {e}")
            return DocumentExtractionResult(
                success=False,
                error_message=f"Word extraction failed: {str(e)}",
                document_type='word'
            )


class DocumentProcessor:
    """High-level processor that integrates document extraction with email parsing."""
    
    def __init__(self, logger: logging.Logger, url_analyzer=None):
        self.logger = logger
        self.document_extractor = DocumentTextExtractor(logger)
        self.url_analyzer = url_analyzer
    
    def process_document_attachment(self, attachment_data: bytes, filename: str = None, 
                                  content_type: str = None) -> Dict[str, Any]:
        """Process a document attachment and extract text and URLs."""
        result = {
            'extraction_result': None,
            'extracted_text': None,
            'urls_found': [],
            'processing_success': False
        }
        
        try:
            # Extract text from document
            extraction_result = self.document_extractor.extract_text(
                attachment_data, filename, content_type
            )
            
            result['extraction_result'] = {
                'success': extraction_result.success,
                'document_type': extraction_result.document_type,
                'extraction_method': extraction_result.extraction_method,
                'error_message': extraction_result.error_message,
                'metadata': extraction_result.metadata
            }
            
            if extraction_result.success and extraction_result.text_content:
                result['extracted_text'] = extraction_result.text_content
                result['processing_success'] = True
                
                # Extract URLs from the document text if URL analyzer available
                if self.url_analyzer:
                    urls_found = self._extract_urls_from_text(extraction_result.text_content)
                    result['urls_found'] = urls_found
                    
                self.logger.info(f"Successfully processed document {filename}: "
                               f"{len(extraction_result.text_content)} chars, "
                               f"{len(result['urls_found'])} URLs")
            else:
                self.logger.warning(f"Failed to extract text from document {filename}: "
                                  f"{extraction_result.error_message}")
        
        except Exception as e:
            self.logger.error(f"Error processing document attachment {filename}: {e}")
            result['extraction_result'] = {
                'success': False,
                'error_message': f"Processing failed: {str(e)}"
            }
        
        return result
    
    def _extract_urls_from_text(self, text: str) -> List[str]:
        """Extract URLs from text content using URL analyzer."""
        try:
            # Create a temporary structure for URL analysis
            temp_structure = {
                'body': {
                    'text': text,
                    'html': None
                }
            }
            
            # Use the URL analyzer to extract URLs
            if hasattr(self.url_analyzer, 'analyze_email_urls'):
                analysis = self.url_analyzer.analyze_email_urls(temp_structure)
                return analysis.final_urls
            else:
                # Fallback to simple regex extraction
                import re
                url_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
                urls = re.findall(url_pattern, text)
                return list(set(urls))  # Remove duplicates
                
        except Exception as e:
            self.logger.error(f"Error extracting URLs from document text: {e}")
            return []